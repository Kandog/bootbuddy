"""
A script to process a Word document template, replacing placeholders in all possible
locations including paragraphs, tables, headers, footers, and text boxes.
"""
import re
import platform
import datetime
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx2pdf import convert

# Regex to find placeholders like v_name, v_date, etc.
PLACEHOLDER_REGEX = re.compile(r'v_[a-zA-Z0-9_]+')

def find_placeholders_in_runs(runs):
    """Extracts placeholder strings from a list of runs."""
    # This is tricky because a placeholder can be split across runs.
    # A simplified approach is to join the text and then find placeholders.
    full_text = "".join(run.text for run in runs)
    return PLACEHOLDER_REGEX.findall(full_text)

def iter_text_runs(element):
    """Yields all text runs in a given element (e.g., paragraph, cell)."""
    if isinstance(element, Paragraph):
        yield from element.runs
    elif hasattr(element, 'paragraphs'):
        for para in element.paragraphs:
            yield from para.runs

def find_placeholders(template_path):
    """
    Finds all unique placeholders in a Word document, searching in paragraphs,
    tables, headers, footers, and text boxes.
    """
    doc = Document(template_path)
    placeholders = set()

    def search_element(element):
        """Finds placeholders in a given element."""
        # Standard paragraphs and tables
        for para in element.paragraphs:
            placeholders.update(find_placeholders_in_runs(para.runs))
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    search_element(cell)

    # Search main body
    search_element(doc)

    # Search headers and footers
    for section in doc.sections:
        if section.header:
            search_element(section.header)
        if section.footer:
            search_element(section.footer)

    # Search text boxes (by parsing the raw XML)
    # This is a more advanced operation, as python-docx has no high-level API for this.
    for el in doc.element.xpath('.//w:txbxContent'):
        for para_element in el.findall('.//' + qn('w:p')):
            p = Paragraph(para_element, doc)
            placeholders.update(find_placeholders_in_runs(p.runs))

    return list(placeholders)


def replace_text_in_runs(runs, data):
    """
    Replaces placeholders in a list of runs.
    This is a simplified implementation. It joins the runs, replaces the text,
    and then overwrites the first run and clears the rest. This will lose
    formatting if the placeholder was split across formatted runs.
    """
    full_text = "".join(run.text for run in runs)

    # Check if any placeholder is present before modifying
    if any(key in full_text for key in data):
        for key, value in data.items():
            if key in full_text:
                full_text = full_text.replace(key, str(value))

        # Overwrite the first run and clear the others
        if runs:
            runs[0].text = full_text
            for run in runs[1:]:
                run.clear()


def replace_placeholders_and_convert(template_path, data, output_docx, output_pdf):
    """
    Replaces placeholders throughout a Word document and converts it to PDF.
    """
    placeholders_in_doc = find_placeholders(template_path)
    missing_keys = [p for p in placeholders_in_doc if p not in data]
    if missing_keys:
        raise ValueError(f"Missing data for placeholders: {', '.join(missing_keys)}")

    doc = Document(template_path)

    def replace_in_element(element):
        """Performs replacement in paragraphs and tables of an element."""
        for para in element.paragraphs:
            replace_text_in_runs(para.runs, data)
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_element(cell)

    # Replace in main body
    replace_in_element(doc)

    # Replace in headers and footers
    for section in doc.sections:
        if section.header:
            replace_in_element(section.header)
        if section.footer:
            replace_in_element(section.footer)

    # Replace in text boxes (by parsing XML)
    # NOTE: This approach has a limitation: it replaces the entire text of a
    # paragraph inside a text box, which may cause loss of complex formatting
    # within that paragraph.
    for el in doc.element.xpath('.//w:txbxContent'):
        for para_element in el.findall('.//' + qn('w:p')):
            p = Paragraph(para_element, doc)
            replace_text_in_runs(p.runs, data)

    # If the last paragraph is empty, remove it to prevent an extra page
    if doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)


    try:
        doc.save(output_docx)
        print(f"Saved modified document to {output_docx}")
        convert(output_docx, output_pdf)
        print(f"Converted {output_docx} to {output_pdf}")
    except Exception as e:
        print(f"Error during PDF conversion: {e}")
        print("Ensure LibreOffice (Linux) or MS Word (Windows) is installed.")


if __name__ == '__main__':
    template_file = 'one_year_template.docx'
    output_docx_file = 'one_year_processed.docx'
    output_pdf_file = 'one_year_processed.pdf'

    print("--- Analyzing Template for All Placeholders (including Text Boxes) ---")
    placeholders = find_placeholders(template_file)
    print("Placeholders found by script:", placeholders)

    full_sample_data = {
        'v_name': 'Johnathan Doe',
        'v_po': 'PO-998877',
        'v_eid': 'EID-112233',
        'v_value': '$275.50',
        'v_epy_date': (datetime.date.today() + datetime.timedelta(days=365)).strftime("%B %d, %Y"),
        'v_issued_date': datetime.date.today().strftime("%B %d, %Y"),
        'v_issued_time': datetime.datetime.now().strftime("%I:%M %p"),
        'v_division': 'Advanced Tech Division',
        'v_cost_centre': 'CC-98765',
        'v_gl': 'GL-54321',
        'v_issued_by': 'Ms. Supervisor'
    }

    data_to_process = {key: full_sample_data[key] for key in placeholders if key in full_sample_data}

    print("\n--- Processing Template ---")
    try:
        replace_placeholders_and_convert(template_file, data_to_process, output_docx_file, output_pdf_file)
        print("\n--- Verifying Output File ---")
        remaining_placeholders = find_placeholders(output_docx_file)
        if not remaining_placeholders:
            print("Verification successful: No placeholders found in the output file.")
        else:
            print("Verification failed: The following placeholders were found in the output file:")
            print(remaining_placeholders)

        print("\nScript finished.")
    except ValueError as e:
        print(f"\nError: {e}")
    except Exception as e:
        print(f"\nAn unexpected error occurred: {e}")
